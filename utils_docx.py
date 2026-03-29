from docx import Document
from docx.shared import Inches
import tempfile
import io

def generar_docx_bytes(template_path, data, graficos, tablas):

    doc = Document(template_path)

    # 🔹 texto
    for p in doc.paragraphs:
        for key, value in data.items():
            if f"{{{{{key}}}}}" in p.text:
                p.text = p.text.replace(f"{{{{{key}}}}}", str(value))

    # 🔹 graficos
    for p in doc.paragraphs:
        for key, fig in graficos.items():
            if f"{{{{{key}}}}}" in p.text:
                p.text = ""

                tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                fig.write_image(tmp.name)

                doc.add_picture(tmp.name, width=Inches(6))

    # 🔹 tablas
    for key, df in tablas.items():
        insertar_tabla(doc, f"{{{{{key}}}}}", df)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer

def insertar_tabla(doc, placeholder, df):

    for p in doc.paragraphs:
        if placeholder in p.text:
            p.text = ""

            table = doc.add_table(rows=1, cols=len(df.columns))

            # 🔹 cabecera
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(df.columns):
                hdr_cells[i].text = str(col)

            # 🔹 filas
            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                for i, val in enumerate(row):
                    row_cells[i].text = str(val)

            return

    