"""
report_generator.py
-------------------
Módulo de generación de informes de optimización de potencias eléctricas.
Genera PDF, HTML y DOCX desde los mismos datos.

Dependencias (solo pip, sin instalaciones en Windows):
    pip install jinja2 xhtml2pdf python-docx matplotlib
"""

import base64
import io
from pathlib import Path
from datetime import datetime

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from jinja2 import Environment, FileSystemLoader
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# ---------------------------------------------------------------------------
# Utilidades
# ---------------------------------------------------------------------------

def fig_to_base64(fig) -> str:
    """Convierte un objeto Figure de matplotlib/plotly a string base64 PNG."""
    buf = io.BytesIO()
    # Soporte matplotlib y plotly
    if hasattr(fig, "savefig"):
        fig.savefig(buf, format="png", bbox_inches="tight", dpi=120)
    elif hasattr(fig, "write_image"):
        fig.write_image(buf, format="png")
    else:
        raise TypeError(f"Tipo de figura no soportado: {type(fig)}")
    buf.seek(0)
    return base64.b64encode(buf.read()).decode("utf-8")


def logo_to_base64(logo_path: str | None) -> str | None:
    """Carga el logo del usuario como base64. Devuelve None si no hay logo."""
    if not logo_path:
        return None
    path = Path(logo_path)
    if not path.exists():
        return None
    with open(path, "rb") as f:
        ext = path.suffix.lower().replace(".", "")
        if ext == "jpg":
            ext = "jpeg"
        data = base64.b64encode(f.read()).decode("utf-8")
        return f"data:image/{ext};base64,{data}"


def fmt_eur(value: float) -> str:
    """Formatea un número como moneda europea."""
    try:
        return f"{float(value):,.2f} €".replace(",", "X").replace(".", ",").replace("X", ".")
    except (TypeError, ValueError):
        return str(value)


# ---------------------------------------------------------------------------
# Preparación del contexto común
# ---------------------------------------------------------------------------

def build_context(
    graf_costes_potcon,
    graf_resumen,
    coste_tp_potcon: float,
    coste_tp_potopt: float,
    ahorro_opt: float,
    ahorro_opt_porc: float,
    df_potencias,
    graf_ahorro,
    graf_costes_pot_periodos,
    logo_path: str | None = None,
    titulo: str = "Informe de Optimización de Potencias",
    subtitulo: str = "",
) -> dict:
    """
    Construye el diccionario de contexto que alimenta tanto la plantilla HTML
    como el generador de DOCX.
    """
    return {
        "titulo": titulo,
        "subtitulo": subtitulo,
        "fecha": datetime.now().strftime("%d/%m/%Y"),
        "logo": logo_to_base64(logo_path),
        # KPIs
        "coste_tp_potcon": fmt_eur(coste_tp_potcon),
        "coste_tp_potopt": fmt_eur(coste_tp_potopt),
        "ahorro_opt": fmt_eur(ahorro_opt),
        "ahorro_opt_porc": f"{float(ahorro_opt_porc):.1f} %",
        # Tabla (HTML desde pandas)
        "tabla_potencias": df_potencias.to_html(
            index=False,
            classes="tabla-datos",
            border=0,
            float_format=lambda x: f"{x:,.2f}",
        ),
        # Gráficos como base64
        "graf_resumen":            fig_to_base64(graf_resumen),
        "graf_costes_potcon":      fig_to_base64(graf_costes_potcon),
        "graf_ahorro":             fig_to_base64(graf_ahorro),
        "graf_costes_pot_periodos": fig_to_base64(graf_costes_pot_periodos),
    }


# ---------------------------------------------------------------------------
# Generador HTML
# ---------------------------------------------------------------------------

def generate_html(context: dict, template_path: str = "templates/informe.html") -> str:
    """
    Renderiza la plantilla Jinja2 con el contexto y devuelve el HTML como string.
    template_path puede ser relativo al directorio de trabajo o absoluto.
    """
    template_dir = str(Path(template_path).parent.resolve())
    template_file = Path(template_path).name
    env = Environment(loader=FileSystemLoader(template_dir))
    template = env.get_template(template_file)
    return template.render(**context)


# ---------------------------------------------------------------------------
# Generador PDF
# ---------------------------------------------------------------------------

def generate_pdf(html_string: str) -> bytes:
    """
    Convierte el HTML a PDF con xhtml2pdf y devuelve los bytes del PDF.
    No requiere instalaciones externas en Windows.
    """
    buf = io.BytesIO()
    pisa_status = pisa.CreatePDF(html_string, dest=buf)
    if pisa_status.err:
        raise RuntimeError(f"Error al generar PDF: {pisa_status.err}")
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Generador DOCX
# ---------------------------------------------------------------------------

def generate_docx(context: dict, df_potencias) -> bytes:
    """
    Genera un archivo Word (.docx) directamente desde los datos (no desde HTML).
    Devuelve los bytes del archivo.
    """
    doc = Document()

    # ---- Estilos generales ----
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ---- Cabecera: logo + título ----
    header = doc.add_heading("", level=0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if context.get("logo"):
        # Decodifica base64 y añade la imagen
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_data = base64.b64decode(context["logo"].split(",")[1])
        logo_buf = io.BytesIO(logo_data)
        run = header_para.add_run()
        run.add_picture(logo_buf, width=Inches(2))

    title_para = doc.add_heading(context["titulo"], level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if context.get("subtitulo"):
        sub = doc.add_paragraph(context["subtitulo"])
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para = doc.add_paragraph(f"Fecha: {context['fecha']}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ---- Resumen KPIs ----
    doc.add_heading("Resumen económico", level=2)
    kpi_table = doc.add_table(rows=2, cols=4)
    kpi_table.style = "Table Grid"
    kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Coste tarifa contratada", "Coste tarifa óptima", "Ahorro estimado", "Ahorro (%)"]
    values  = [
        context["coste_tp_potcon"],
        context["coste_tp_potopt"],
        context["ahorro_opt"],
        context["ahorro_opt_porc"],
    ]
    accent = RGBColor(0x1A, 0x56, 0xDB)  # azul corporativo

    for i, (h, v) in enumerate(zip(headers, values)):
        hcell = kpi_table.cell(0, i)
        hcell.text = h
        hcell.paragraphs[0].runs[0].bold = True
        hcell.paragraphs[0].runs[0].font.color.rgb = accent
        kpi_table.cell(1, i).text = v

    doc.add_paragraph()

    # ---- Tabla de potencias ----
    doc.add_heading("Detalle de potencias", level=2)
    cols = list(df_potencias.columns)
    data_table = doc.add_table(rows=1 + len(df_potencias), cols=len(cols))
    data_table.style = "Table Grid"
    data_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, col in enumerate(cols):
        cell = data_table.cell(0, i)
        cell.text = col
        cell.paragraphs[0].runs[0].bold = True

    for row_idx, row in df_potencias.iterrows():
        for col_idx, val in enumerate(row):
            data_table.cell(row_idx + 1, col_idx).text = str(val)

    doc.add_paragraph()

    # ---- Gráficos ----
    doc.add_heading("Análisis gráfico", level=2)
    graficos = [
        ("Resumen",                   context["graf_resumen"]),
        ("Costes potencia contratada", context["graf_costes_potcon"]),
        ("Ahorro estimado",            context["graf_ahorro"]),
        ("Costes por periodos",        context["graf_costes_pot_periodos"]),
    ]
    for nombre, b64 in graficos:
        doc.add_paragraph(nombre).runs[0].bold = True
        img_data = base64.b64decode(b64)
        img_buf = io.BytesIO(img_data)
        doc.add_picture(img_buf, width=Inches(5.5))
        doc.add_paragraph()

    # ---- Bytes de salida ----
    out_buf = io.BytesIO()
    doc.save(out_buf)
    out_buf.seek(0)
    return out_buf.read()


# ---------------------------------------------------------------------------
# Función principal: genera los tres formatos de una vez
# ---------------------------------------------------------------------------

def generar_informe(
    graf_costes_potcon,
    graf_resumen,
    coste_tp_potcon: float,
    coste_tp_potopt: float,
    ahorro_opt: float,
    ahorro_opt_porc: float,
    df_potencias,
    graf_ahorro,
    graf_costes_pot_periodos,
    logo_path: str | None = None,
    titulo: str = "Informe de Optimización de Potencias",
    subtitulo: str = "",
    template_path: str = "templates/informe.html",
) -> dict:
    """
    Genera los tres formatos del informe y devuelve un dict con las claves:
        - "html"  → str  (HTML completo)
        - "pdf"   → bytes
        - "docx"  → bytes

    Uso en Streamlit:
        resultado = generar_informe(...)
        st.download_button("Descargar PDF",  resultado["pdf"],  "informe.pdf",  "application/pdf")
        st.download_button("Descargar Word", resultado["docx"], "informe.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        st.download_button("Descargar HTML", resultado["html"].encode(), "informe.html", "text/html")
    """
    context = build_context(
        graf_costes_potcon=graf_costes_potcon,
        graf_resumen=graf_resumen,
        coste_tp_potcon=coste_tp_potcon,
        coste_tp_potopt=coste_tp_potopt,
        ahorro_opt=ahorro_opt,
        ahorro_opt_porc=ahorro_opt_porc,
        df_potencias=df_potencias,
        graf_ahorro=graf_ahorro,
        graf_costes_pot_periodos=graf_costes_pot_periodos,
        logo_path=logo_path,
        titulo=titulo,
        subtitulo=subtitulo,
    )

    html_str  = generate_html(context, template_path)
    pdf_bytes = generate_pdf(html_str)
    docx_bytes = generate_docx(context, df_potencias)

    return {"html": html_str, "pdf": pdf_bytes, "docx": docx_bytes}
